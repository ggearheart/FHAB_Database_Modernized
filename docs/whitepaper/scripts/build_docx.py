#!/usr/bin/env python3
"""Regenerate the editable Word draft of the FHAB lab-data linkage white paper.

Output:  ../FHAB-Lab-Data-Linkage-Whitepaper-DRAFT.docx
Usage:   python3 scripts/build_docx.py        (run from docs/whitepaper/)
Requires: python-docx  (pip install --user python-docx)

This script is the single source of truth for the .docx. It mirrors index.html; when the page
content changes in a way that should reach the download, edit the text here and re-run — do not
hand-edit the .docx. Palette mirrors the page (deep teal / HAB green / SEC amber).
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---- palette (mirrors index.html) ----
TEAL = RGBColor(0x0D, 0x3B, 0x3A)   # brand / headings
GREEN = RGBColor(0x2F, 0x7D, 0x5B)  # accent
AMBER = RGBColor(0xB4, 0x53, 0x0A)  # SEC / eyebrow
INK = RGBColor(0x18, 0x28, 0x27)
BODY = RGBColor(0x2C, 0x3E, 0x3C)
SUB = RGBColor(0x4A, 0x5C, 0x5A)
TAG_BG = "12302F"
PAPER = "F3F6F5"
LINE = "DBE4E1"

OUT = os.path.join(os.path.dirname(__file__), "..",
                   "FHAB-Lab-Data-Linkage-Whitepaper-DRAFT.docx")

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = BODY
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.28

sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.9)
sec.left_margin = sec.right_margin = Inches(1.05)

for lvl, sz in (("Heading 1", 17), ("Heading 2", 13.5), ("Heading 3", 11)):
    st = doc.styles[lvl]
    st.font.name = "Georgia"
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = TEAL if lvl != "Heading 3" else INK
    st.paragraph_format.space_before = Pt(14 if lvl == "Heading 1" else 10)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True


# ---- helpers ----
def _shade(el, fill):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), fill)
    el.append(sh)


def rule(color="DBE4E1", size=6, space=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space); p.paragraph_format.space_after = Pt(space + 5)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), color)
    pbdr.append(bottom); pPr.append(pbdr)
    return p


def eyebrow(text, color=GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper()); r.font.name = "Consolas"; r.font.size = Pt(7.5)
    r.font.color.rgb = color; r.font.bold = True
    rPr = r._element.get_or_add_rPr(); sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), "40"); rPr.append(sp)
    return p


def body(text, italic=False, color=BODY, after=8):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    _add_runs(p, text, italic=italic, color=color)
    return p


def _add_runs(p, text, italic=False, color=BODY):
    """Split on **bold** and *italic* markers."""
    import re
    for chunk in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text):
        if not chunk:
            continue
        r = p.add_run()
        if chunk.startswith("**") and chunk.endswith("**"):
            r.text = chunk[2:-2]; r.bold = True; r.font.color.rgb = INK
        elif chunk.startswith("*") and chunk.endswith("*"):
            r.text = chunk[1:-1]; r.italic = True; r.font.color.rgb = color
        else:
            r.text = chunk; r.italic = italic; r.font.color.rgb = color


def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    _add_runs(p, text)
    return p


def section(num, kicker, title):
    eyebrow(f"§ {num} — {kicker}", AMBER)
    doc.add_heading(title, level=2)


# ============================ CONTENT ============================
eyebrow("FHAB Database Modernization · Technical White Paper · Draft", GREEN)
h = doc.add_heading("Linking Laboratory Data to a Bloom Incident — Before the Samples Are Even Taken",
                    level=1)
deck = doc.add_paragraph()
deck.paragraph_format.space_after = Pt(6)
r = deck.add_run("A sampling-provenance approach: issue the identifier at the moment of sampling, and "
                 "let event → sample → result linkage happen by construction, not by archaeology.")
r.font.name = "Georgia"; r.italic = True; r.font.size = Pt(11.5); r.font.color.rgb = SUB
meta = doc.add_paragraph()
mr = meta.add_run("Prepared by the FHAB Database Modernization project   ·   Audience: FHAB program & "
                  "Regional Water Board staff   ·   Status: Draft for discussion, July 2026")
mr.font.name = "Consolas"; mr.font.size = Pt(8); mr.font.color.rgb = RGBColor(0x7C, 0x8C, 0x8A)
rule(color="0D3B3A", size=12)

body("**The hardest part of managing HAB laboratory data isn't the chemistry — it's the timing.** An "
     "incident opens before anyone knows who will sample, how many samples they'll collect, or which "
     "analyses they'll run. Samples flow to one or more laboratories; results return a week or more "
     "later; high-concern cases spawn additional sampling events over weeks or months. Each artifact "
     "lands in a folder, and the work of connecting it back to the right report, case, and response is "
     "deferred — sometimes for years. This paper argues that the linkage problem is fundamentally a "
     "problem of *when identity is assigned*, and proposes a single, low-friction fix: a **Sampling "
     "Event Code** minted at the first hint of sampling and carried on every communication that follows.")

section("01", "The problem", "A moving timeline that outruns its own recordkeeping")
body("Freshwater HAB response is not a tidy pipeline. A typical incident unfolds like this:")
bullet("An incident is **opened on partial information** — a report, a photo, a phone call. Whether "
       "anyone samples, and who, is often undecided.")
bullet("One or more parties — Water Board staff, a tribe, a lake association, a partner agency — "
       "**decide independently to sample**, each on their own timeline.")
bullet("Samples are collected, split, and shipped to **different laboratories**, each with its own "
       "submission forms, sample IDs, and turnaround.")
bullet("Results **trickle back about a week later** — as spreadsheets, PDFs, and chain-of-custody "
       "scans, under the lab's naming, not the program's.")
bullet("For a high-concern case, **additional sampling events follow**, repeating the cycle and "
       "multiplying the artifacts.")
body("Everything gets “tucked into folders.” Months or years later, someone opens those folders "
     "and tries to reconstruct which result belongs to which sampling event, which event to which case, "
     "and which case to which response and advisory. That reconstruction is slow, ambiguous, and lossy "
     "— and it delays the moment the data can become authoritative public information.")

section("02", "Root cause", "We assign identity too late")
body("From a database perspective, the difficulty has one root: **the shared key that ties an analytical "
     "result to its incident does not exist at the time the sample is taken.** It has to be "
     "*reconstructed later* from weak, circumstantial signals:")
bullet("**Fuzzy matching after the fact.** Identity is inferred from station code, sample date, "
       "coordinates, and water-body name. The modernized system already does this — a station "
       "registry plus a “within ~2 km and ±14 days” reconciler — but that is a "
       "*heuristic backstop*, not ground truth. Two blooms on the same lake in the same fortnight are "
       "genuinely hard to tell apart from coordinates alone.")
bullet("**The Chain-of-Custody carries the only authoritative location and date** — and it arrives "
       "as paper or a rotated PDF that must be read and transcribed by hand, often long after the fact.")
bullet("**Every party uses its own identifiers.** The field crew, each lab, and the case file name the "
       "same sample three different ways, with no key in common.")
bullet("**One sampling event is fragmented across records** — grab samples to Lab A, SPATT passive "
       "samplers to Lab B, genetic assays to Lab C — with nothing declaring they belong together.")
body("No amount of downstream cleverness fully recovers information that was never captured. The fix is "
     "to move identity *upstream*, to the one moment every party passes through: the act of sampling.")

section("03", "The proposal", "One code, issued at sampling time, on every communication")
body("Issue a **Sampling Event Code (SEC)** the moment a sampling effort is first anticipated — "
     "*before* you know the party, the sample count, or the analyses. The SEC is a short, human-usable, "
     "collision-free token that goes on **everything**: the sampling tasking or order, the "
     "chain-of-custody, the sample labels, the lab submission, the returned results file, the folder and "
     "file names, and the case record.")

# SEC label block (shaded table)
t = doc.add_table(rows=2, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = True
cell0 = t.cell(0, 0).merge(t.cell(0, 3))
for c in [cell0] + [t.cell(1, i) for i in range(4)]:
    _shade(c._tc.get_or_add_tcPr(), TAG_BG)
p = cell0.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
rk = p.add_run("SAMPLING EVENT CODE  ·  AFFIX TO CoC & EVERY DELIVERABLE")
rk.font.name = "Consolas"; rk.font.size = Pt(7); rk.font.bold = True
rk.font.color.rgb = RGBColor(0xF0, 0xA3, 0x4A)
pc = cell0.add_paragraph(); rc = pc.add_run("FHAB-R5-20260731-7QK2")
rc.font.name = "Consolas"; rc.font.size = Pt(16); rc.font.bold = True
rc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for i, (code, lab) in enumerate([("FHAB", "Program"), ("R5", "Regional board"),
                                 ("20260731", "Event opened"), ("7QK2", "Random + check")]):
    cc = t.cell(1, i)
    pp = cc.paragraphs[0]; rr = pp.add_run(code)
    rr.font.name = "Consolas"; rr.font.bold = True; rr.font.size = Pt(9.5)
    rr.font.color.rgb = RGBColor(0xF0, 0xA3, 0x4A)
    pl = cc.add_paragraph(); rl = pl.add_run(lab)
    rl.font.size = Pt(7.5); rl.font.color.rgb = RGBColor(0xCF, 0xDA, 0xD3)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

body("The format is deliberately **self-describing** so a person can sanity-check it at a glance, with a "
     "trailing check character so a mistyped code is caught at entry rather than silently mis-linked. "
     "Crucially, the SEC is **scope-free at issue**: the responsible party, the number of samples, the "
     "target analytes, and the destination labs all attach to the code *as they materialize*. You never "
     "have to know the shape of the sampling effort in advance — you only have to name it.")
doc.add_heading("It maps onto the model you already have", level=3)
body("The SEC is the human-facing face of an identifier the database already uses. The modernized schema "
     "represents each ingested batch as a **Sampling Event** with its own Sampling_Event_ID, drawn from a "
     "reserved id range (≥ 1,000,000,000) that can never collide with the integer keys published on "
     "data.ca.gov. The SEC simply **pre-registers** that sampling event and pushes its identifier upstream "
     "to the field and the lab, so the batch exists *before* its data arrives instead of being conjured "
     "at ingest. The relationships are already in place:")
pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = pr.add_run("Case  →  Report / Event  →  Sampling Event (SEC)  →  Sample  →  Result")
rr.font.name = "Consolas"; rr.font.size = Pt(10); rr.font.color.rgb = TEAL; rr.bold = True
body("A single high-concern case naturally holds **many** sampling events over time — each new field "
     "effort mints a new SEC under the same case, which is exactly the “additional sampling for the "
     "cases we worry about” pattern.")

section("04", "The workflow", "Pre-register, then reconcile on arrival")
body("The SEC inverts today's order of operations. Instead of “collect data, then reconstruct "
     "identity,” the flow becomes “declare identity, then collect data against it.”")
for t_, head, txt in [
    ("T0 · sampling anticipated", "Open a sampling event",
     "A staffer clicks “Start a sampling event.” The system mints the SEC and produces a "
     "pre-stamped field packet — CoC header, sample labels, and a QR code — carrying the SEC, "
     "the region, and the water body if known. Party and analyses are left blank."),
    ("Field", "Everything wears the code",
     "Samples are labeled with the SEC; the CoC is pre-filled with it. Splits sent to different labs all "
     "carry the same SEC plus a short sample suffix, so the fragments stay a family."),
    ("Laboratory", "The code rides the submission and returns",
     "The SEC travels on the lab submission and is required back on the electronic deliverable — a "
     "dedicated field, or a filename convention. This is a one-line ask in a lab contract or SOP."),
    ("≈ 1 week later · results arrive", "Linking is a lookup, not a guess",
     "Ingest matches on the SEC and auto-attaches every sample and result to the pre-registered sampling "
     "event — and thus to the case. Results that arrive without a code fall to the existing "
     "workboard reconciler as the exception path, not the default."),
]:
    eyebrow(t_, AMBER)
    doc.add_heading(head, level=3)
    body(txt, after=6)
q = doc.add_paragraph(); q.paragraph_format.left_indent = Inches(0.25)
q.paragraph_format.space_before = Pt(6)
pPr = q._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); left = OxmlElement("w:left")
left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "24"); left.set(qn("w:space"), "8")
left.set(qn("w:color"), "B4530A"); pbdr.append(left); pPr.append(pbdr)
rq = q.add_run("Move the identifier to the front of the timeline, and the reconciliation work that "
               "consumes staff months mostly disappears — because it never has to happen.")
rq.font.name = "Georgia"; rq.italic = True; rq.font.size = Pt(13); rq.font.color.rgb = TEAL

section("05", "Build vs. buy-forward", "Most of the machinery already exists")
body("This is an increment on the modernized system, not a rebuild. The linking spine, reconciler, and "
     "provenance controls are in place; what's new is pushing the identifier upstream.")
rows = [
    ("Capability", "Status", "Role in the SEC workflow"),
    ("Sampling Event batch + Sampling_Event_ID", "In place", "Becomes the record the SEC pre-registers"),
    ("Reserved id range (≥ 1e9)", "In place", "Mints SECs that never collide with published ids"),
    ("Station registry (shared codes)", "In place", "Common location vocabulary for field & labs"),
    ("Workboard reconciler (2 km / ±14 d)", "In place", "Fallback for uncoded / legacy data"),
    ("Duplicate detect & merge", "In place", "Trivial once an SEC is present"),
    ("Provenance & row-level audit", "In place", "Every SEC-linked ingest is attributable"),
    ("Pre-issued SEC + printable field packet", "To add", "The front door of the workflow"),
    ("Lab EDD template + “return the code” clause", "To add", "Guarantees the code comes home"),
    ("SEC auto-match on ingest", "To add", "Turns linking into a lookup"),
    ("QR labels / mobile CoC capture", "Later", "Removes hand-transcription entirely"),
]
tbl = doc.add_table(rows=len(rows), cols=3); tbl.style = "Light List Accent 1"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, txt in enumerate(rows[0]):
    c = tbl.cell(0, ci); c.paragraphs[0].runs and None
    run = c.paragraphs[0].add_run(txt) if not c.paragraphs[0].runs else c.paragraphs[0].runs[0]
    run.text = txt; run.font.name = "Consolas"; run.font.size = Pt(7.5); run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _shade(c._tc.get_or_add_tcPr(), "0D3B3A")
for ri, row in enumerate(rows[1:], start=1):
    for ci, txt in enumerate(row):
        c = tbl.cell(ri, ci); p = c.paragraphs[0]
        r = p.add_run(txt); r.font.size = Pt(8.5)
        if ci == 1:
            r.bold = True
            r.font.color.rgb = GREEN if txt == "In place" else AMBER
        else:
            r.font.color.rgb = BODY

section("06", "Supporting practices", "Habits that make the code stick")
bullet("**One code, every communication.** The SEC belongs on the CoC, the sample order, the lab "
       "deliverable, and the folder and file names — plus a scannable QR/barcode for one-tap linking.")
bullet("**Standardize the deliverable.** Ask labs for a CEDEN-aligned electronic data deliverable "
       "carrying the SEC and station code, instead of a PDF to be transcribed. Provide a template; make "
       "it the default in agreements.")
bullet("**Capture the CoC as data at T0.** A mobile chain-of-custody form stamped with the SEC makes the "
       "authoritative location and date digital from the first minute — ending the "
       "rotated-PDF-and-retype problem.")
bullet("**Share the station registry.** When field crews and labs draw station codes from the same "
       "published list, locations line up by construction.")
bullet("**Keep the reconciler forever.** Legacy folders and the occasional un-coded submission will "
       "always exist; the heuristic matcher remains the safety net beneath the code.")

section("07", "Why it's worth it", "What good looks like")
kpis = [("months → same-week", "Time to link a result to its case"),
        ("> 90%", "Results auto-linked on arrival (target)"),
        ("↓ duplicates", "One code makes doubles obvious"),
        ("↑ speed", "Faster path to authoritative public data")]
kt = doc.add_table(rows=1, cols=4); kt.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (big, lbl) in enumerate(kpis):
    c = kt.cell(0, i); _shade(c._tc.get_or_add_tcPr(), PAPER)
    p = c.paragraphs[0]; r = p.add_run(big)
    r.font.name = "Georgia"; r.bold = True; r.font.size = Pt(13); r.font.color.rgb = TEAL
    pl = c.add_paragraph(); rl = pl.add_run(lbl); rl.font.size = Pt(8); rl.font.color.rgb = SUB
doc.add_paragraph().paragraph_format.space_after = Pt(2)
body("Beyond the numbers: complete provenance for every result, far less staff time spent on folder "
     "archaeology, and the confidence that a determination or advisory rests on evidence that is "
     "unambiguously tied to its incident.")

section("08", "Getting there", "A phased, low-risk rollout")
for n, head, txt in [
    ("Phase 1", "Mint & match (now)",
     "Generate the SEC in the app, print it on a CoC cover sheet and sample labels, ask labs to return "
     "it, and auto-link on an SEC match at ingest. Uses the sampling-event batch and reconciler that "
     "already exist — no field hardware required."),
    ("Phase 2", "Scan & standardize",
     "Add QR labels and a mobile CoC form for digital capture at T0, and publish a CEDEN-aligned lab "
     "deliverable template so results arrive machine-ready."),
    ("Phase 3", "Direct submission",
     "Offer partners and labs a portal or API to submit deliverables keyed by SEC, closing the loop from "
     "field to public data with no manual hop."),
]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    rn = p.add_run(n + "   "); rn.font.name = "Consolas"; rn.bold = True; rn.font.size = Pt(9)
    rn.font.color.rgb = AMBER
    rh = p.add_run(head); rh.bold = True; rh.font.color.rgb = INK; rh.font.size = Pt(10.5)
    body(txt, after=8)

section("09", "Honest limits", "Risks & how we blunt them")
bullet("**Adoption depends on people using the code.** Mitigate by making the SEC the path of least "
       "resistance — pre-printed packets, a one-click issue button, and a “return the "
       "code” clause in lab agreements — with the reconciler catching whatever slips through.")
bullet("**Parties who don't coordinate.** Whoever opens the incident can issue and share the SEC; where "
       "several parties sample independently, each event gets its own SEC and all of them reconcile "
       "cleanly under one case.")
bullet("**Legacy data.** Everything already in folders keeps flowing through the existing heuristic "
       "matcher; the SEC changes the future, not the past.")

section("10", "Recommendation", "Adopt the Sampling Event Code as the linking spine")
body("Pre-register every sampling event and push its code upstream to the field and the labs; "
     "standardize the electronic deliverable so the code comes home; and keep heuristic reconciliation "
     "as the exception path, not the default. The database is already built for it — the SEC simply "
     "lets us capture, at the one moment every party shares, the small piece of information that today we "
     "spend months trying to recover.")

rule(color="0D3B3A", size=12)
f = doc.add_paragraph()
rf = f.add_run("Draft for discussion. ")
rf.bold = True; rf.font.color.rgb = INK; rf.font.size = Pt(9)
rf2 = f.add_run("Prepared from the FHAB Database Modernization project's data-governance perspective as "
                "input to program and Regional Water Board planning. Terminology (report / event / case / "
                "response / advisory, Sampling Event, station registry) follows the modernized FHAB schema "
                "and the published CA FHAB data model.")
rf2.font.size = Pt(9); rf2.font.color.rgb = SUB

doc.save(OUT)
print(f"Wrote {os.path.abspath(OUT)}")
